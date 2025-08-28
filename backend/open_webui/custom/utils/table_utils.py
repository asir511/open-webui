# utils/table_utils.py
# -*- coding: utf-8 -*-
from __future__ import annotations
import io, os, tempfile
import re
from typing import List, Dict, Any, Optional

# 尽量可选依赖，按可用性回退
try:
    import camelot  # 需要 ghostscript & opencv
    _HAS_CAMELOT = True
except Exception:
    _HAS_CAMELOT = False

try:
    import tabula  # 需要 JRE
    _HAS_TABULA = True
except Exception:
    _HAS_TABULA = False

try:
    import pdfplumber
    _HAS_PDFPLUMBER = True
except Exception:
    _HAS_PDFPLUMBER = False

try:
    import docx  # python-docx
    _HAS_DOXC = True
except Exception:
    _HAS_DOXC = False

import pandas as pd


# ========== 基础渲染：DataFrame -> HTML / GFM ==========

def df_to_html(df: pd.DataFrame, title: Optional[str] = None) -> str:
    """渲染一个 DataFrame 为独立 HTML 片段（自带极简样式）"""
    style = """
    <style>
      .tbl {border-collapse: collapse; width: 100%; margin: 8px 0;}
      .tbl th, .tbl td {border: 1px solid #ddd; padding: 6px; text-align: left; white-space: pre-wrap;}
      .tbl thead th {background: #f7f7f8; font-weight: 600;}
      .tbl caption {text-align: left; font-size: 13px; color: #666; padding: 4px 0;}
      .tbl tr:nth-child(even) {background: #fafafa;}
    </style>
    """
    cap = f"<caption>{escape_html(title)}</caption>" if title else ""
    html_table = df.to_html(classes="tbl", index=False, border=0, escape=False)
    # pandas 会把 <table ...> 包好，我们插入 caption
    html_table = html_table.replace("<table ", f"<table {''} ").replace("</table>", "")
    # 简单插 caption（如果有）
    if cap:
        html_table = html_table.replace("<thead>", f"{cap}<thead>")
    html_table += "</table>"
    return style + html_table


def df_to_gfm(df: pd.DataFrame) -> str:
    """渲染一个 DataFrame 为 GitHub-Flavored Markdown 管道表格（不依赖 tabulate）"""
    # 所有值转字符串，避免 None 报错
    cols = [str(c) if c is not None else "" for c in df.columns.tolist()]
    rows = [[_safe_str(v) for v in row] for row in df.astype(object).values.tolist()]
    # 计算每列宽度
    widths = [len(c) for c in cols]
    for r in rows:
        for j, v in enumerate(r):
            widths[j] = max(widths[j], len(v))
    # 拼 header
    head = "| " + " | ".join(c.ljust(widths[i]) for i, c in enumerate(cols)) + " |"
    sep  = "| " + " | ".join("-" * widths[i] for i, _ in enumerate(cols)) + " |"
    # 拼 rows
    body_lines = []
    for r in rows:
        body_lines.append("| " + " | ".join(r[j].ljust(widths[j]) for j in range(len(widths))) + " |")
    return "\n".join([head, sep] + body_lines)


def escape_html(s: Optional[str]) -> str:
    if s is None:
        return ""
    return (s.replace("&", "&amp;")
             .replace("<", "&lt;")
             .replace(">", "&gt;")
             .replace('"', "&quot;")
             .replace("'", "&#39;"))


def _safe_str(v: Any) -> str:
    if v is None:
        return ""
    s = str(v)
    # 把换行替换成 <br>，避免破坏 GFM 的单元格
    return re.sub(r"\r?\n", "<br>", s)


# ========== PDF 表格提取 ==========

def extract_pdf_tables_to_dfs(pdf_bytes: bytes) -> List[pd.DataFrame]:
    """
    从 PDF 提取表格并返回 DataFrame 列表。
    优先 Camelot(lattice/stream) → Tabula → pdfplumber。
    可能为空列表（无表或依赖缺失）。
    """
    dfs: List[pd.DataFrame] = []

    # 1) Camelot（需要依赖完整）
    if _HAS_CAMELOT:
        try:
            with io.BytesIO(pdf_bytes) as bio:
                # Camelot 不接内存流，只接路径；这里绕不过去，只能写临时文件在调用方做。
                pass
        except Exception:
            pass

    # 2) Tabula
    if _HAS_TABULA and not dfs:
        try:
            with io.BytesIO(pdf_bytes) as bio:
                # tabula 也需要路径，无法直接 bytes。由上层落地后传路径更好。
                pass
        except Exception:
            pass

    # 3) pdfplumber（最通用，纯 Python）
    if _HAS_PDFPLUMBER and not dfs:
        try:
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                for page in pdf.pages:
                    tables = page.extract_tables()
                    for t in tables or []:
                        if not t or not any(any(cell for cell in row) for row in t):
                            continue
                        # 第一行当 header
                        header = [(c or "").strip() for c in (t[0] or [])]
                        data_rows = [
                            [(c or "").strip() for c in row]
                            for row in (t[1:] or [])
                        ]
                        # 某些表头会重复为空，简单兜底一下
                        ncols = max(len(header), *(len(r) for r in data_rows)) if data_rows else len(header)
                        header += [""] * (ncols - len(header))
                        norm_rows = [r + [""] * (ncols - len(r)) for r in data_rows]
                        df = pd.DataFrame(norm_rows, columns=header[:ncols])
                        # 丢弃完全空的列
                        df = df.loc[:, (df != "").any(axis=0)]
                        if df.shape[1] == 0:
                            continue
                        dfs.append(df)
        except Exception:
            pass

    return dfs


# ========== DOCX 表格提取 ==========

def extract_docx_tables_to_dfs(docx_bytes: bytes) -> List[pd.DataFrame]:
    """
    用 python-docx 提取 .docx 表格
    """
    if not _HAS_DOXC:
        return []
    from docx import Document
    doc = Document(io.BytesIO(docx_bytes))
    dfs: List[pd.DataFrame] = []
    for tbl in doc.tables:
        rows = []
        for row in tbl.rows:
            cells = []
            for cell in row.cells:
                # 合并单元格在 python-docx 层面较难恢复，这里先取最终文本
                txt = cell.text.replace("\t", " ").strip()
                rows.append([_safe_str(c.text if hasattr(c, "text") else txt) for c in row.cells])
                break  # 注意：上面已经收集本行，这里 break 避免重复
        # 上述写法会多收；改为逐行收集:
    dfs = []
    for tbl in doc.tables:
        tmp = []
        for row in tbl.rows:
            tmp.append([_safe_str(cell.text) for cell in row.cells])
        if not tmp:
            continue
        # 第一行作 header
        header = tmp[0]
        data = tmp[1:]
        # 填齐列长
        ncols = max(len(header), *(len(r) for r in data)) if data else len(header)
        header += [""] * (ncols - len(header))
        data = [r + [""] * (ncols - len(r)) for r in data]
        df = pd.DataFrame(data, columns=header[:ncols])
        df = df.loc[:, (df != "").any(axis=0)]
        if df.shape[1]:
            dfs.append(df)
    return dfs


# ========== 汇总导出：HTML + GFM ==========

def dfs_pack_as_render(dfs: List[pd.DataFrame], caption_prefix: str = "Table") -> Dict[str, Any]:
    """
    输入 DataFrame 列表，输出：
      - tables_meta: [{index, rows, cols}]
      - tables_html: 合并的 HTML 字符串（含多个表）
      - tables_gfm: 合并的 GFM 文本（含多个表）
    """
    html_parts: List[str] = []
    gfm_parts: List[str] = []
    metas: List[Dict[str, Any]] = []

    for i, df in enumerate(dfs, 1):
        title = f"{caption_prefix} {i} ({df.shape[0]}x{df.shape[1]})"
        metas.append({"index": i, "rows": int(df.shape[0]), "cols": int(df.shape[1])})
        html_parts.append(df_to_html(df, title=title))
        gfm_parts.append(f"**{title}**\n\n{df_to_gfm(df)}")

    return {
        "tables_meta": metas,
        "tables_html": "\n\n".join(html_parts) if html_parts else "",
        "tables_gfm": "\n\n\n".join(gfm_parts) if gfm_parts else "",
    }

def extract_pdf_tables_with_camelot_path(
    pdf_path: str,
    pages: str = "all",
    flavors: Optional[list[str]] = None,
    strip_text: str = "\n",           # 去掉换行提升单元格合并稳定性
    line_scale: float = 40,           # lattice 识别线的尺度；可调 15~50
    shift_text: List[int] | None = None,  # stream 模式文本偏移
    table_areas: Optional[list[str]] = None,  # 指定识别区域 ["x1,y1,x2,y2"]，可调试
    columns: Optional[list[str]] = None,      # 指定列竖线坐标
    dpi: int = 200
) -> List[pd.DataFrame]:
    """
    用 Camelot 从 PDF 路径提表，返回 DataFrame 列表。
    flavors: ['lattice','stream'] 的顺序决定优先级；不传则两种都跑一遍去重拼接。
    """
    try:
        import camelot
    except Exception as e:
        raise RuntimeError("缺少依赖：camelot。请先安装 camelot-py[cv]、ghostscript、opencv。") from e

    if flavors is None:
        flavors = ["lattice", "stream"]

    dfs: List[pd.DataFrame] = []
    seen_csv: set[str] = set()

    for flavor in flavors:
        try:
            tables = camelot.read_pdf(
                pdf_path,
                pages=pages,
                flavor=flavor,
                strip_text=strip_text,
                line_scale=line_scale if flavor == "lattice" else None,
                shift_text=shift_text if flavor == "stream" else None,
                table_areas=table_areas,
                columns=columns,
                dpi=dpi,
            )
        except Exception as e:
            # 某些页/文件在某 flavor 下会报错，直接跳过尝试下一个
            continue

        for t in tables or []:
            df = t.df if getattr(t, "df", None) is not None else None
            if df is None or df.shape[1] == 0:
                continue
            # 规范化：去掉全空列
            df = df.astype(str).fillna("").replace("^nan$", "", regex=True)
            df = df.loc[:, (df != "").any(axis=0)]
            if df.shape[1] == 0:
                continue
            # 去重：以 CSV 文本作为近似指纹
            csv_sig = df.to_csv(index=False)
            if csv_sig in seen_csv:
                continue
            seen_csv.add(csv_sig)
            dfs.append(df)

    return dfs


def extract_pdf_tables_with_camelot_bytes(
    pdf_bytes: bytes,
    **kwargs
) -> List[pd.DataFrame]:
    """
    把 bytes 落临时文件，再走 Camelot 提取。
    kwargs 透传给 extract_pdf_tables_with_camelot_path
    """
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=True) as tmp:
        tmp.write(pdf_bytes)
        tmp.flush()
        return extract_pdf_tables_with_camelot_path(tmp.name, **kwargs)